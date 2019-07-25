from docx import Document
from docx.shared import Inches

from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

def makeDOC(fName, directory, companyName):
    os.chdir(directory)

    document = Document()

    document.add_heading(companyName + " Social Media Posts")
    document.add_paragraph("\n\n")

    iterator = 1
    with open(fName, "r") as fp:
        for line in fp:
            picName = str(iterator) + ".png"
            document.add_picture(picName, width = Inches(4))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER 

            paragraph = document.add_paragraph(line)
            

            document.add_page_break()

            iterator += 1

    paragraph = document.paragraphs[-1]
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None
    document.save(companyName + "-SM-approval.docx")

